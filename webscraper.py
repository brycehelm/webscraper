import sys
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import time
from collections import deque
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Configure logging
def setup_logging(base_domain):
    log_file = f"{base_domain}_crawl.log"
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file),
            logging.StreamHandler()
        ]
    )
    return log_file

def crawl(start_url, base_domain, max_pages=1000, initial_batch_size=5, keywords=None):
    def is_valid_url(url, visited_set):
        parsed = urlparse(url)
        # Skip common irrelevant paths
        path = parsed.path.lower()
        irrelevant_patterns = [
            '/changelog', '/updates', '/archive', '/log', 
            '/privacy', '/terms', '/sitemap', '/feed',
            '/wp-content', '/wp-includes', '/tag/', '/category/',
            '/author/', '/comment', '/trackback', '/feed', '/rss'
        ]
        
        if any(pattern in path for pattern in irrelevant_patterns):
            return False
            
        return parsed.netloc == base_domain and url not in visited_set

    # Initialize with default keywords if none provided
    if keywords is None:
        keywords = []
    
    queue = deque([start_url])
    visited = set()
    discovered_urls = set([start_url])
    content_list = []  # To store scraped data
    url_scores = {start_url: 1.0}  # Relevance scores for URLs

    logging.info(f"Starting crawl from {start_url} with max_pages={max_pages}")

    try:
        while queue and len(discovered_urls) < max_pages:
            # Sort queue by relevance score if we have enough URLs
            if len(queue) > 10:
                # Convert to list, sort by score (descending), and back to deque
                url_list = list(queue)
                url_list.sort(key=lambda u: url_scores.get(u, 0.0), reverse=True)
                queue = deque(url_list)
            
            url = queue.popleft()
            if url in visited:
                logging.debug(f"Skipping already visited URL: {url}")
                continue
            visited.add(url)

            logging.info(f"Fetching URL: {url} (Discovered: {len(discovered_urls)})")
            try:
                response = requests.get(url, timeout=5)
                response.raise_for_status()

                if 'text/html' not in response.headers.get('Content-Type', ''):
                    logging.warning(f"Skipping non-HTML content at {url}")
                    continue

                soup = BeautifulSoup(response.text, 'html.parser')
                title = soup.title.string if soup.title else url
                title = title.strip()

                # Extract and clean text
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text(separator=" ").strip()
                lines = (line.strip() for line in text.splitlines())
                text = " ".join(line for line in lines if line)
                
                # Calculate relevance score based on content length and keywords
                relevance_score = min(1.0, len(text) / 5000)  # Length-based score
                
                # Boost score based on keywords if provided
                if keywords:
                    text_lower = text.lower()
                    keyword_matches = sum(keyword.lower() in text_lower for keyword in keywords)
                    keyword_score = min(1.0, keyword_matches / len(keywords))
                    relevance_score = 0.7 * relevance_score + 0.3 * keyword_score
                
                # Skip pages with very low relevance
                if relevance_score < 0.2:
                    logging.info(f"Skipping low relevance page: {url} (score: {relevance_score:.2f})")
                    continue

                # Store the data
                content_list.append({
                    'url': url,
                    'title': title,
                    'text': text,
                    'relevance': relevance_score
                })

                links_found = 0
                for link in soup.find_all('a', href=True):
                    absolute_url = urljoin(url, link['href'])
                    if is_valid_url(absolute_url, visited):
                        if absolute_url not in discovered_urls:
                            links_found += 1
                            # Calculate initial score for new URL based on link text and position
                            link_text = link.get_text().strip().lower()
                            link_score = relevance_score * 0.8  # Inherit some relevance from parent
                            
                            # Boost score if link text contains keywords
                            if keywords and any(keyword.lower() in link_text for keyword in keywords):
                                link_score += 0.2
                                
                            url_scores[absolute_url] = link_score
                            
                        discovered_urls.add(absolute_url)
                        if absolute_url not in visited:
                            queue.append(absolute_url)
                logging.debug(f"Found {links_found} new links on {url}")

            except requests.RequestException as e:
                logging.error(f"Failed to fetch {url}: {e}")
                content_list.append({
                    'url': url,
                    'title': 'Error',
                    'text': f"Failed to fetch content - {str(e)}",
                    'relevance': 0.0
                })

            time.sleep(0.1)

    except KeyboardInterrupt:
        logging.info("KeyboardInterrupt detected. Saving collected content...")
        raise

    # Sort content by relevance before returning
    content_list.sort(key=lambda x: x.get('relevance', 0.0), reverse=True)
    logging.info(f"Crawl completed. Processed {len(content_list)} pages.")
    return content_list

def save_to_docx(content_list, base_domain):
    doc = Document()

    # Title page
    title = doc.add_paragraph(f"Web Scraping Report for {base_domain}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True

    subtitle = doc.add_paragraph(f"Generated on {time.strftime('%B %d, %Y')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # Content
    for entry in content_list:
        # Add header (page title or URL)
        header_text = entry['title'] if entry['title'] != 'Error' else entry['url']
        doc.add_heading(header_text, level=1)

        # Add body text
        body = doc.add_paragraph(entry['text'])
        body.style.font.size = Pt(12)
        body.style.font.name = 'Times New Roman'
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

    # Always include timestamp in filename to avoid conflicts
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    output_file = f"{base_domain}_report_{timestamp}.docx"
    
    try:
        doc.save(output_file)
        
        # Add the output file to .gitignore
        try:
            gitignore_path = r"C:\Users\brych\OneDrive\Documents\Python Scripts\WebSraper\.gitignore"
            
            # Create the file if it doesn't exist
            import os
            if not os.path.exists(gitignore_path):
                with open(gitignore_path, 'w') as f:
                    f.write("# Generated report files\n")
            
            with open(gitignore_path, 'a+') as gitignore:
                # First check if the file is already in .gitignore
                gitignore.seek(0)
                content = gitignore.read()
                
                # Add patterns for all report files if not already present
                patterns_to_add = []
                if f"{base_domain}_report_*.docx" not in content:
                    patterns_to_add.append(f"{base_domain}_report_*.docx")
                if "*.docx" not in content:
                    patterns_to_add.append("*.docx")
                
                # Write the new patterns
                if patterns_to_add:
                    gitignore.write('\n' + '\n'.join(patterns_to_add) + '\n')
                    logging.info(f"Added report patterns to .gitignore")
        except Exception as e:
            logging.warning(f"Could not update .gitignore: {e}")
            
    except Exception as e:
        # If saving fails, try the user's home directory
        import os
        home_dir = os.path.expanduser("~")
        output_file = os.path.join(home_dir, f"{base_domain}_report_{timestamp}.docx")
        try:
            doc.save(output_file)
        except Exception as e:
            logging.error(f"Failed to save report: {e}")
            return None
    
    logging.info(f"Saved report to {output_file}")
    return output_file

def main():
    if len(sys.argv) < 2:
        print("Usage: python webscraper.py <url> [keyword1 keyword2 ...]")
        print("Example: python webscraper.py example.com product service")
        sys.exit(1)

    input_url = sys.argv[1]
    if not input_url.startswith(('http://', 'https://')):
        start_url = f"https://{input_url}"
    else:
        start_url = input_url

    # Extract optional keywords from command line
    keywords = sys.argv[2:] if len(sys.argv) > 2 else []
    
    base_domain = urlparse(start_url).netloc
    log_file = setup_logging(base_domain)

    try:
        content_list = crawl(
            start_url=start_url,
            base_domain=base_domain,
            max_pages=1000,
            initial_batch_size=5,
            keywords=keywords
        )
    except KeyboardInterrupt:
        print("\nStopped by user.")
    else:
        print(f"\nCrawling complete. Processed {len(content_list)} pages.")

    # Save to Word document
    output_file = save_to_docx(content_list, base_domain)
    print(f"Report saved to {output_file}")
    print(f"Log saved to {log_file}")
    
    if keywords:
        print(f"Used keywords for relevance: {', '.join(keywords)}")

if __name__ == "__main__":
    main()